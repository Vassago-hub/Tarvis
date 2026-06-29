from __future__ import annotations

from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "submissions" / "track-one-tcl-service-agent" / "TECHNICAL_DOCUMENT.md"
OUTPUT = ROOT / "submissions" / "track-one-tcl-service-agent" / "TCL官网智能服务助手_技术文档.docx"

FONT = "Microsoft YaHei"
CODE_FONT = "Microsoft YaHei"


def set_run_font(run, size: float | None = None, bold: bool | None = None, color: str | None = None, font: str = FONT):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run._element.rPr.rFonts.set(qn("w:ascii"), font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_style_font(style, size: float, bold: bool = False, color: str = "000000", font: str = FONT):
    style.font.name = font
    style._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    style._element.rPr.rFonts.set(qn("w:ascii"), font)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), font)
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = RGBColor.from_string(color)


def set_paragraph_spacing(style, before: float = 0, after: float = 6, line: float = 1.1):
    pf = style.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line


def shade_paragraph(paragraph, fill: str):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), "9360")


def add_inline_markdown(paragraph, text: str, base_size: float = 11):
    pattern = re.compile(r"(`[^`]+`|\*\*[^*]+\*\*)")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos:match.start()])
            set_run_font(run, base_size)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, base_size, bold=True)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, base_size, font=CODE_FONT, color="1F2937")
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(run, base_size)


def build_doc() -> None:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    set_style_font(styles["Normal"], 11, False, "000000")
    set_paragraph_spacing(styles["Normal"], 0, 6, 1.1)
    set_style_font(styles["Heading 1"], 16, True, "2E74B5")
    set_paragraph_spacing(styles["Heading 1"], 16, 8, 1.1)
    set_style_font(styles["Heading 2"], 13, True, "2E74B5")
    set_paragraph_spacing(styles["Heading 2"], 12, 6, 1.1)
    set_style_font(styles["Heading 3"], 12, True, "1F4D78")
    set_paragraph_spacing(styles["Heading 3"], 8, 4, 1.1)
    set_style_font(styles["Title"], 20, True, "0B2545")
    set_paragraph_spacing(styles["Title"], 0, 12, 1.1)
    set_style_font(styles["List Bullet"], 11, False, "000000")
    set_style_font(styles["List Number"], 11, False, "000000")

    markdown = SOURCE.read_text(encoding="utf-8")
    lines = markdown.splitlines()
    in_code = False
    code_lines: list[str] = []
    table_lines: list[str] = []

    def flush_code():
        nonlocal code_lines
        if not code_lines:
            return
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(7)
        p.paragraph_format.left_indent = Inches(0.12)
        p.paragraph_format.right_indent = Inches(0.08)
        shade_paragraph(p, "F8FAFC")
        run = p.add_run("\n".join(code_lines))
        set_run_font(run, 9, font=CODE_FONT, color="111827")
        code_lines = []

    def flush_table():
        nonlocal table_lines
        if not table_lines:
            return
        rows = []
        for line in table_lines:
            cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
            if all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells):
                continue
            rows.append(cells)
        if rows:
            table = doc.add_table(rows=len(rows), cols=len(rows[0]))
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.LEFT
            table.autofit = False
            set_table_width(table)
            for row_idx, row in enumerate(rows):
                for col_idx, value in enumerate(row):
                    cell = table.cell(row_idx, col_idx)
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    set_cell_margins(cell)
                    if row_idx == 0:
                        set_cell_shading(cell, "F2F4F7")
                    paragraph = cell.paragraphs[0]
                    paragraph.paragraph_format.space_after = Pt(0)
                    add_inline_markdown(paragraph, value, 9.5)
                    for run in paragraph.runs:
                        if row_idx == 0:
                            run.bold = True
            doc.add_paragraph()
        table_lines = []

    for line in lines:
        if line.startswith("```"):
            if in_code:
                flush_code()
                in_code = False
            else:
                flush_table()
                in_code = True
            continue
        if in_code:
            code_lines.append(line)
            continue
        if line.startswith("|"):
            table_lines.append(line)
            continue
        flush_table()
        text = line.strip()
        if not text:
            continue
        if text.startswith("# "):
            p = doc.add_paragraph(style="Title")
            add_inline_markdown(p, text[2:], 20)
        elif text.startswith("## "):
            p = doc.add_paragraph(style="Heading 1")
            add_inline_markdown(p, text[3:], 16)
        elif text.startswith("### "):
            p = doc.add_paragraph(style="Heading 2")
            add_inline_markdown(p, text[4:], 13)
        elif text.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_inline_markdown(p, text[2:], 11)
        else:
            p = doc.add_paragraph()
            add_inline_markdown(p, text, 11)

    flush_table()
    flush_code()

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("TCL 官网智能服务助手技术文档")
    set_run_font(run, 8, color="666666")

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_doc()

