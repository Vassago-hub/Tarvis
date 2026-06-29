from __future__ import annotations

from pathlib import Path
import re

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "submissions" / "track-one-tcl-service-agent" / "TARVIS_PET_DESIGN.md"
OUTPUT = ROOT / "submissions" / "track-one-tcl-service-agent" / "Tarvis宠物精灵设计文档.docx"
FONT = "Microsoft YaHei"


def set_run_font(run, size=None, bold=None, color=None, font=FONT):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run._element.rPr.rFonts.set(qn("w:ascii"), font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_style(style, size, bold=False, color="000000"):
    style.font.name = FONT
    style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = RGBColor.from_string(color)


def add_inline(paragraph, text, size=11):
    pattern = re.compile(r"(`[^`]+`|\*\*[^*]+\*\*)")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos:match.start()])
            set_run_font(run, size)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size, bold=True)
        else:
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, size, color="1F2937")
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(run, size)


def shade_paragraph(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, val in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")


def build_doc():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    set_style(styles["Normal"], 11)
    styles["Normal"].paragraph_format.space_after = Pt(6)
    styles["Normal"].paragraph_format.line_spacing = 1.1
    set_style(styles["Title"], 20, True, "0B2545")
    set_style(styles["Heading 1"], 16, True, "2E74B5")
    set_style(styles["Heading 2"], 13, True, "2E74B5")
    set_style(styles["Heading 3"], 12, True, "1F4D78")
    set_style(styles["List Bullet"], 11)

    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    in_code = False
    code_lines = []
    table_lines = []

    def flush_code():
        nonlocal code_lines
        if not code_lines:
            return
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.12)
        p.paragraph_format.space_after = Pt(7)
        shade_paragraph(p, "F8FAFC")
        run = p.add_run("\n".join(code_lines))
        set_run_font(run, 9, color="111827")
        code_lines = []

    def flush_table():
        nonlocal table_lines
        if not table_lines:
            return
        rows = []
        for line in table_lines:
            cells = [c.strip() for c in line.strip().strip("|").split("|")]
            if all(re.fullmatch(r":?-{3,}:?", c) for c in cells):
                continue
            rows.append(cells)
        if rows:
            table = doc.add_table(rows=len(rows), cols=len(rows[0]))
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.LEFT
            table.autofit = True
            for ri, row in enumerate(rows):
                for ci, value in enumerate(row):
                    cell = table.cell(ri, ci)
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    cell_margins(cell)
                    if ri == 0:
                        shade_cell(cell, "F2F4F7")
                    p = cell.paragraphs[0]
                    p.paragraph_format.space_after = Pt(0)
                    add_inline(p, value, 9.5)
                    for run in p.runs:
                        if ri == 0:
                            run.bold = True
            doc.add_paragraph()
        table_lines = []

    for line in lines:
        if line.startswith("```"):
            if in_code:
                flush_code()
                in_code = False
            else:
                flush_table()
                in_code = True
            continue
        if in_code:
            code_lines.append(line)
            continue
        if line.startswith("|"):
            table_lines.append(line)
            continue
        flush_table()
        text = line.strip()
        if not text:
            continue
        if text.startswith("# "):
            p = doc.add_paragraph(style="Title")
            add_inline(p, text[2:], 20)
        elif text.startswith("## "):
            p = doc.add_paragraph(style="Heading 1")
            add_inline(p, text[3:], 16)
        elif text.startswith("### "):
            p = doc.add_paragraph(style="Heading 2")
            add_inline(p, text[4:], 13)
        elif text.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, text[2:], 11)
        else:
            p = doc.add_paragraph()
            add_inline(p, text, 11)

    flush_table()
    flush_code()

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("Tarvis 宠物精灵设计文档")
    set_run_font(run, 8, color="666666")

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_doc()
